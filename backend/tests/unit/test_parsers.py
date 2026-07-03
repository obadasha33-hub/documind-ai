"""Unit tests for document parsers."""

import pytest
from io import BytesIO

from app.services.parsers import parse_pdf, parse_docx, parse_markdown, ParsedPage


class TestPDFParser:
    """Tests for PDF parsing."""

    def test_parse_valid_pdf(self, sample_pdf_content):
        """Test parsing a valid PDF."""
        result = parse_pdf(BytesIO(sample_pdf_content))
        assert result is not None
        assert len(result) > 0
        assert "Hello World" in result[0].content

    def test_parse_empty_pdf(self):
        """Test parsing an empty PDF raises error."""
        with pytest.raises(Exception):  # pymupdf.EmptyFileError
            parse_pdf(BytesIO(b""))

    def test_parse_invalid_pdf(self):
        """Test parsing invalid PDF raises error."""
        with pytest.raises(Exception):  # pymupdf.FileDataError
            parse_pdf(BytesIO(b"not a pdf"))


class TestDOCXParser:
    """Tests for DOCX parsing."""

    def test_parse_valid_docx(self):
        """Test parsing a valid DOCX."""
        from docx import Document
        doc = Document()
        doc.add_paragraph("Test content")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        result = parse_docx(buffer.getvalue())
        assert len(result) > 0
        assert "Test content" in result[0].content

    def test_parse_empty_docx(self):
        """Test parsing empty DOCX."""
        from docx import Document
        doc = Document()
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        result = parse_docx(buffer.getvalue())
        assert result == []


class TestMarkdownParser:
    """Tests for Markdown parsing."""

    def test_parse_valid_markdown(self, sample_markdown_content):
        """Test parsing valid markdown."""
        result = parse_markdown(sample_markdown_content)
        assert len(result) > 0
        assert "# Sample Document" in result[0].content
        assert "DocuMind AI" in result[0].content

    def test_parse_empty_markdown(self):
        """Test parsing empty markdown."""
        result = parse_markdown(b"")
        assert result == []

    def test_parse_markdown_with_code_blocks(self):
        """Test parsing markdown with code blocks."""
        content = b"""# Test
        
```python
print("hello")
```
"""
        result = parse_markdown(content)
        assert len(result) > 0
        assert "print" in result[0].content