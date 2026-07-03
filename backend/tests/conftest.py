"""Pytest configuration and fixtures."""

import pytest
import pytest_asyncio
import uuid
import httpx
from httpx import AsyncClient, ASGITransport
from unittest.mock import AsyncMock, MagicMock
from sqlalchemy.ext.asyncio import create_async_engine, AsyncSession, async_sessionmaker
from sqlalchemy.pool import StaticPool

from app.main import app
from app.core.database import get_db, Base
from app.core.security import get_password_hash, create_access_token
from app.models.user import User, UserRole
from app.models.tenant import Tenant, Plan
from app.models.document import Document, DocumentStatus
from app.models.conversation import Conversation
from app.models.message import Message, MessageRole


# Test database URL (in-memory SQLite for fast tests)
TEST_DATABASE_URL = "sqlite+aiosqlite:///:memory:"


@pytest.fixture(scope="session")
def event_loop():
    """Create event loop for async tests."""
    import asyncio
    loop = asyncio.get_event_loop_policy().new_event_loop()
    yield loop
    loop.close()


@pytest_asyncio.fixture(scope="function")
async def db_engine():
    """Create test database engine."""
    engine = create_async_engine(
        TEST_DATABASE_URL,
        connect_args={"check_same_thread": False},
        poolclass=StaticPool,
    )
    async with engine.begin() as conn:
        await conn.run_sync(Base.metadata.create_all)
    yield engine
    await engine.dispose()


@pytest_asyncio.fixture(scope="function")
async def db_session(db_engine):
    """Create test database session."""
    async_session = async_sessionmaker(db_engine, class_=AsyncSession, expire_on_commit=False)
    async with async_session() as session:
        yield session


@pytest_asyncio.fixture(scope="function")
async def async_client(db_session):
    """Create async test client with overridden database."""
    async def override_get_db():
        yield db_session
    
    app.dependency_overrides[get_db] = override_get_db
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        yield client
    app.dependency_overrides.clear()


@pytest.fixture
def sample_pdf_content() -> bytes:
    """Return minimal valid PDF content for testing."""
    return b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 44 >>
stream
BT /F1 12 Tf 100 700 Td (Hello World) Tj ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f
0000000009 00000 n
0000000058 00000 n
0000000115 00000 n
0000000266 00000 n
0000000360 00000 n
trailer
<< /Size 6 /Root 1 0 R >>
startxref
441
%%EOF"""


@pytest.fixture
def sample_markdown_content() -> bytes:
    """Return sample markdown content."""
    return b"""# Sample Document

This is a test document for DocuMind AI.

## Section 1
The quick brown fox jumps over the lazy dog. This sentence contains all letters of the alphabet.

## Section 2
DocuMind AI is a multi-tenant RAG platform that helps teams query their documents using natural language.

## Section 3
The platform supports PDF, DOCX, and Markdown file formats.
"""


@pytest.fixture
def sample_docx_content() -> bytes:
    """Return sample DOCX content."""
    from docx import Document
    from io import BytesIO
    doc = Document()
    doc.add_heading("Test Document", 0)
    doc.add_paragraph("This is a test document for DocuMind AI.")
    doc.add_heading("Section 1", 1)
    doc.add_paragraph("Content for section 1.")
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


@pytest_asyncio.fixture
async def test_tenant(db_session):
    """Create a test tenant."""
    tenant = Tenant(
        id=str(uuid.uuid4()),
        name="Test Tenant",
        slug="test-tenant",
        plan=Plan.FREE,
    )
    db_session.add(tenant)
    await db_session.commit()
    await db_session.refresh(tenant)
    return tenant


@pytest_asyncio.fixture
async def test_user(db_session, test_tenant):
    """Create a test user."""
    user = User(
        id=str(uuid.uuid4()),
        email="test@example.com",
        name="Test User",
        role=UserRole.ADMIN,
        tenant_id=test_tenant.id,
    )
    db_session.add(user)
    await db_session.commit()
    await db_session.refresh(user)
    return user


@pytest.fixture
def auth_headers(test_user):
    """Create authentication headers for test user."""
    token = create_access_token(data={"sub": str(test_user.id)})
    return {"Authorization": f"Bearer {token}"}


@pytest_asyncio.fixture
async def test_documents(db_session, test_tenant, test_user):
    """Create test documents."""
    docs = [
        Document(
            id=str(uuid.uuid4()),
            tenant_id=test_tenant.id,
            filename=f"doc_{i}.pdf",
            file_type="pdf",
            file_size=1024,
            r2_key=f"{test_tenant.id}/{uuid.uuid4()}/doc_{i}.pdf",
            status=DocumentStatus.READY,
            chunk_count=5,
        )
        for i in range(3)
    ]
    db_session.add_all(docs)
    await db_session.commit()
    for doc in docs:
        await db_session.refresh(doc)
    return docs


@pytest_asyncio.fixture
async def test_document(db_session, test_tenant, test_user):
    """Create a single test document."""
    doc = Document(
        id=str(uuid.uuid4()),
        tenant_id=test_tenant.id,
        filename="test.pdf",
        file_type="pdf",
        file_size=1024,
        r2_key=f"{test_tenant.id}/{uuid.uuid4()}/test.pdf",
        status=DocumentStatus.READY,
        chunk_count=5,
    )
    db_session.add(doc)
    await db_session.commit()
    await db_session.refresh(doc)
    return doc


@pytest_asyncio.fixture
async def test_chat_sessions(db_session, test_tenant, test_user):
    """Create test chat sessions."""
    sessions = [
        Conversation(
            id=str(uuid.uuid4()),
            tenant_id=test_tenant.id,
            user_id=test_user.id,
            title=f"Session {i}"
        )
        for i in range(2)
    ]
    db_session.add_all(sessions)
    await db_session.commit()
    for s in sessions:
        await db_session.refresh(s)
    return sessions


@pytest_asyncio.fixture
async def test_chat_session(db_session, test_tenant, test_user):
    """Create a single test chat session."""
    session = Conversation(
        id=str(uuid.uuid4()),
        tenant_id=test_tenant.id,
        user_id=test_user.id,
        title="Test Session"
    )
    db_session.add(session)
    await db_session.commit()
    await db_session.refresh(session)
    return session


@pytest_asyncio.fixture
async def test_chat_messages(db_session, test_chat_session, test_user):
    """Create test chat messages."""
    messages = [
        Message(
            id=str(uuid.uuid4()),
            conversation_id=test_chat_session.id,
            role=MessageRole.USER,
            content="Test question"
        ),
        Message(
            id=str(uuid.uuid4()),
            conversation_id=test_chat_session.id,
            role=MessageRole.ASSISTANT,
            content="Test answer",
            sources=[{
                "chunk_id": "1",
                "document_id": "1",
                "document_name": "test.pdf",
                "content": "citation",
                "score": 0.9
            }]
        )
    ]
    db_session.add_all(messages)
    await db_session.commit()
    return messages


@pytest.fixture
def mock_gemini_client():
    """Mock Gemini client for testing."""
    mock = MagicMock()
    mock.embed_content = AsyncMock()
    mock.generate_content = AsyncMock()
    return mock


@pytest.fixture
def mock_groq_client():
    """Mock Groq client for testing."""
    mock = MagicMock()
    mock.chat.completions.create = AsyncMock()
    return mock


@pytest.fixture
def mock_redis():
    """Mock Redis client for testing."""
    mock = AsyncMock()
    mock.get = AsyncMock(return_value=None)
    mock.set = AsyncMock()
    mock.delete = AsyncMock()
    mock.incr = AsyncMock(return_value=1)
    return mock